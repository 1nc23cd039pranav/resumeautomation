from docx import Document


def generate_tailored_resume(base_resume_path, output_path, jd_skills, ats_score):
    doc = Document(base_resume_path)

    doc.add_paragraph("\nATS Optimization Summary")
    doc.add_paragraph(f"ATS Match Score: {ats_score}%")
    doc.add_paragraph("Targeted Skills: " + ", ".join(jd_skills))

    doc.save(output_path)